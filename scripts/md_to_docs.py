"""Convert EXPANSION_PLAN.md to .docx and .pdf with the same content.

Handles the markdown subset used in the doc: ATX headings (#..####), bullet and
numbered lists, GFM tables, blockquotes, fenced code blocks, horizontal rules,
**bold**, and `inline code`. Not a general markdown engine — tuned for this doc.

Run:  python scripts/md_to_docs.py
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import (
    HRFlowable,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

ROOT = Path(__file__).resolve().parent.parent
SRC = ROOT / "docs" / "EXPANSION_PLAN.md"
DOCX_OUT = ROOT / "docs" / "EXPANSION_PLAN.docx"
PDF_OUT = ROOT / "docs" / "EXPANSION_PLAN.pdf"

ACCENT = RGBColor(0x4F, 0x46, 0xE5)  # indigo, matches the app


# ---------------------------------------------------------------------------
# Parse markdown into a flat list of block dicts.
# ---------------------------------------------------------------------------
def parse(md: str) -> list[dict]:
    blocks: list[dict] = []
    lines = md.splitlines()
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # Fenced code block
        if stripped.startswith("```"):
            i += 1
            code: list[str] = []
            while i < n and not lines[i].strip().startswith("```"):
                code.append(lines[i])
                i += 1
            i += 1  # skip closing fence
            blocks.append({"type": "code", "text": "\n".join(code)})
            continue

        # Blank line
        if not stripped:
            i += 1
            continue

        # Horizontal rule
        if re.fullmatch(r"-{3,}", stripped):
            blocks.append({"type": "hr"})
            i += 1
            continue

        # Heading
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            blocks.append({"type": "h", "level": len(m.group(1)), "text": m.group(2)})
            i += 1
            continue

        # Table (header row followed by a |---| separator)
        if stripped.startswith("|") and i + 1 < n and re.match(r"^\s*\|?[\s:|-]+\|", lines[i + 1]):
            rows: list[list[str]] = []
            while i < n and lines[i].strip().startswith("|"):
                rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
                i += 1
            # rows[1] is the --- separator; drop it
            header = rows[0]
            body = rows[2:] if len(rows) > 2 else []
            blocks.append({"type": "table", "header": header, "rows": body})
            continue

        # Blockquote
        if stripped.startswith(">"):
            quote: list[str] = []
            while i < n and lines[i].strip().startswith(">"):
                quote.append(lines[i].strip().lstrip(">").strip())
                i += 1
            blocks.append({"type": "quote", "text": " ".join(quote)})
            continue

        # Numbered list item
        m = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if m:
            blocks.append({"type": "li", "ordered": True, "text": m.group(2)})
            i += 1
            continue

        # Bullet list item
        if re.match(r"^[-*]\s+", stripped):
            blocks.append({"type": "li", "ordered": False, "text": re.sub(r"^[-*]\s+", "", stripped)})
            i += 1
            continue

        # Paragraph (gather consecutive plain lines)
        para: list[str] = [stripped]
        i += 1
        while i < n and lines[i].strip() and not re.match(
            r"^(#{1,6}\s|>|[-*]\s|\d+\.\s|\||```|-{3,}$)", lines[i].strip()
        ):
            para.append(lines[i].strip())
            i += 1
        blocks.append({"type": "p", "text": " ".join(para)})
    return blocks


def strip_inline(text: str) -> str:
    """Remove markdown link syntax/backticks for plain rendering, keep link text."""
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)  # [t](url) -> t
    return text


# ---------------------------------------------------------------------------
# DOCX renderer
# ---------------------------------------------------------------------------
def render_docx(blocks: list[dict]) -> None:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    def add_runs(p, text: str) -> None:
        # Split on **bold** and `code`, render runs accordingly.
        text = strip_inline(text)
        for part in re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text):
            if not part:
                continue
            if part.startswith("**") and part.endswith("**"):
                r = p.add_run(part[2:-2]); r.bold = True
            elif part.startswith("`") and part.endswith("`"):
                r = p.add_run(part[1:-1]); r.font.name = "Consolas"
            else:
                p.add_run(part)

    for b in blocks:
        t = b["type"]
        if t == "h":
            lvl = b["level"]
            if lvl == 1:
                p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                r = p.add_run(strip_inline(b["text"])); r.bold = True; r.font.size = Pt(22)
                r.font.color.rgb = ACCENT
            else:
                h = doc.add_heading(level=min(lvl, 4))
                h.runs.clear() if h.runs else None
                r = h.add_run(strip_inline(b["text"]))
                r.font.color.rgb = ACCENT
        elif t == "p":
            add_runs(doc.add_paragraph(), b["text"])
        elif t == "li":
            style = "List Number" if b["ordered"] else "List Bullet"
            add_runs(doc.add_paragraph(style=style), b["text"])
        elif t == "quote":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            r = p.add_run(strip_inline(b["text"])); r.italic = True
            r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        elif t == "code":
            p = doc.add_paragraph()
            r = p.add_run(b["text"]); r.font.name = "Consolas"; r.font.size = Pt(9)
        elif t == "hr":
            doc.add_paragraph("_" * 50)
        elif t == "table":
            cols = len(b["header"])
            tbl = doc.add_table(rows=1, cols=cols)
            tbl.style = "Light Grid Accent 1"
            for j, cell in enumerate(b["header"]):
                run = tbl.rows[0].cells[j].paragraphs[0].add_run(strip_inline(cell))
                run.bold = True
            for row in b["rows"]:
                cells = tbl.add_row().cells
                for j in range(cols):
                    val = row[j] if j < len(row) else ""
                    add_runs(cells[j].paragraphs[0], val)
    doc.save(str(DOCX_OUT))
    print(f"DOCX -> {DOCX_OUT}")


# ---------------------------------------------------------------------------
# PDF renderer (reportlab)
# ---------------------------------------------------------------------------
def _rl_inline(text: str) -> str:
    text = strip_inline(text)
    text = (text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"))
    text = re.sub(r"\*\*([^*]+)\*\*", r"<b>\1</b>", text)
    text = re.sub(r"`([^`]+)`", r'<font face="Courier">\1</font>', text)
    return text


def render_pdf(blocks: list[dict]) -> None:
    styles = getSampleStyleSheet()
    accent = colors.HexColor("#4F46E5")
    h1 = ParagraphStyle("h1x", parent=styles["Title"], textColor=accent, fontSize=22, spaceAfter=10, alignment=TA_LEFT)
    h2 = ParagraphStyle("h2x", parent=styles["Heading2"], textColor=accent, fontSize=15, spaceBefore=12)
    h3 = ParagraphStyle("h3x", parent=styles["Heading3"], textColor=accent, fontSize=12, spaceBefore=8)
    body = ParagraphStyle("bodyx", parent=styles["BodyText"], fontSize=9.5, leading=13, spaceAfter=5)
    bullet = ParagraphStyle("bulletx", parent=body, leftIndent=14, bulletIndent=4)
    quote = ParagraphStyle("quotex", parent=body, leftIndent=12, textColor=colors.HexColor("#555555"), fontName="Helvetica-Oblique", backColor=colors.HexColor("#F3F2FF"), borderPadding=4)
    code = ParagraphStyle("codex", parent=body, fontName="Courier", fontSize=8, backColor=colors.HexColor("#F4F4F4"), leftIndent=6, borderPadding=4)

    story: list = []
    for b in blocks:
        t = b["type"]
        if t == "h":
            style = {1: h1, 2: h2, 3: h3}.get(b["level"], h3)
            story.append(Paragraph(_rl_inline(b["text"]), style))
        elif t == "p":
            story.append(Paragraph(_rl_inline(b["text"]), body))
        elif t == "li":
            prefix = "• "
            story.append(Paragraph(prefix + _rl_inline(b["text"]), bullet))
        elif t == "quote":
            story.append(Paragraph(_rl_inline(b["text"]), quote))
        elif t == "code":
            safe = b["text"].replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace("\n", "<br/>")
            story.append(Paragraph(safe, code))
            story.append(Spacer(1, 4))
        elif t == "hr":
            story.append(Spacer(1, 4))
            story.append(HRFlowable(width="100%", color=colors.HexColor("#DDDDDD")))
            story.append(Spacer(1, 4))
        elif t == "table":
            cols = len(b["header"])
            data = [[Paragraph(_rl_inline(c), ParagraphStyle("th", parent=body, textColor=colors.white, fontName="Helvetica-Bold", fontSize=9)) for c in b["header"]]]
            for row in b["rows"]:
                data.append([Paragraph(_rl_inline(row[j] if j < len(row) else ""), ParagraphStyle("td", parent=body, fontSize=8.5)) for j in range(cols)])
            avail = A4[0] - 30 * mm
            tbl = Table(data, colWidths=[avail / cols] * cols, repeatRows=1)
            tbl.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), accent),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#CCCCCC")),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F7F7FB")]),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ]))
            story.append(tbl)
            story.append(Spacer(1, 6))

    SimpleDocTemplate(
        str(PDF_OUT), pagesize=A4,
        leftMargin=15 * mm, rightMargin=15 * mm, topMargin=15 * mm, bottomMargin=15 * mm,
        title="Intants — Production Expansion Plan",
    ).build(story)
    print(f"PDF  -> {PDF_OUT}")


if __name__ == "__main__":
    md = SRC.read_text(encoding="utf-8")
    blocks = parse(md)
    render_docx(blocks)
    render_pdf(blocks)
    print("Done.")
